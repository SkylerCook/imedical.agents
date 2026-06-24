#!/usr/bin/env python3
"""Parse interface documents into Markdown and structured artifacts."""

from __future__ import annotations

import argparse
import json
import re
import shutil
import subprocess
import sys
import zipfile
from xml.etree import ElementTree as ET
from dataclasses import dataclass, asdict, field as dataclass_field
from pathlib import Path
from typing import Any, Iterable


CODE_HEADERS = {"\u53c2\u6570\u540d", "\u53c2\u6570\u540d\u79f0", "\u53c2\u6570\u4ee3\u7801", "\u53c2\u6570", "\u5b57\u6bb5", "\u5b57\u6bb5\u540d", "\u5b57\u6bb5\u540d\u79f0", "\u5b57\u6bb5\u4ee3\u7801", "\u5b57\u6bb5\u7f16\u7801", "\u6570\u636e\u9879\u4ee3\u7801", "\u6570\u636e\u5b57\u6bb5\u540d", "\u6570\u636e\u5143\u82f1\u6587\u540d\u79f0", "\u4ee3\u7801", "field", "field name", "code", "parameter", "param"}
NAME_HEADERS = {"\u63cf\u8ff0", "\u4e2d\u6587\u540d", "\u53c2\u6570\u540d\u79f0", "\u6570\u636e\u9879\u540d\u79f0", "\u5b57\u6bb5\u63cf\u8ff0", "\u6570\u636e\u5143\u4e2d\u6587\u540d\u79f0", "\u540d\u79f0", "name"}
TYPE_HEADERS = {"\u6570\u636e\u7c7b\u578b", "\u6570\u636e\u9879\u7c7b\u578b", "\u53c2\u6570\u7c7b\u578b", "\u7c7b\u578b", "\u5b57\u6bb5\u7c7b\u578b", "type", "datatype", "data type"}
LENGTH_HEADERS = {"\u957f\u5ea6", "\u6570\u636e\u957f\u5ea6", "\u5b57\u6bb5\u957f\u5ea6", "length", "len"}
REQUIRED_HEADERS = {"\u662f\u5426\u5fc5\u586b", "\u5fc5\u586b", "\u662f\u5426\u5fc5\u987b", "\u975e\u7a7a", "required", "mandatory", "not null"}
NULLABLE_HEADERS = {"\u5141\u8bb8\u7a7a", "\u53ef\u7a7a", "\u662f\u5426\u5141\u8bb8\u7a7a", "nullable", "allow null"}
PRIMARY_KEY_HEADERS = {"\u4e3b\u952e", "\u662f\u5426\u4e3b\u952e", "primary key", "pk"}
DEFAULT_HEADERS = {"\u9ed8\u8ba4\u503c", "\u7f3a\u7701\u503c", "default", "default value"}
DESC_HEADERS = {"\u5907\u6ce8", "\u8bf4\u660e", "\u5b57\u6bb5\u8bf4\u660e", "\u63cf\u8ff0", "description", "remark", "remarks"}

@dataclass
class Field:
    code: str
    name: str
    fieldType: str = ""
    length: str = ""
    required: str = ""
    nullable: str = ""
    primaryKey: str = ""
    defaultValue: str = ""
    description: str = ""
    jsonPath: str = ""
    requiredByMarker: bool = False
    requiredMismatch: bool = False
    sourceHeaderMap: dict[str, str] | None = None
    rawColumns: dict[str, str] = dataclass_field(default_factory=dict)
    sourceLocation: dict[str, Any] = dataclass_field(default_factory=dict)
    classification: str = "mapped-field"
    confidence: float = 1.0
    warnings: list[str] = dataclass_field(default_factory=list)
    requiredReason: str = ""
    jsonPathReason: str = ""

@dataclass
class View:
    viewCode: str
    viewName: str
    fields: list[Field]



@dataclass
class PdfContext:
    interfaceNumber: str = ""
    interfaceTitle: str = ""
    subsectionNumber: str = ""
    subsectionTitle: str = ""
    parameterObject: str = ""


def normalize_header(value: Any) -> str:
    text = str(value or "").strip()
    return re.sub(r"\s+", " ", text).lower()


def cell_text(value: Any) -> str:
    if value is None:
        return ""
    text = str(value).strip()
    return re.sub(r"\s+", " ", text)


def slugify(path: Path) -> str:
    stem = path.stem.strip().lower()
    stem = re.sub(r"[^a-z0-9\u4e00-\u9fff._-]+", "-", stem)
    stem = re.sub(r"-+", "-", stem).strip("-._")
    return stem or "interface-document"


def markdown_table(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    header = normalized[0]
    lines = [
        "| " + " | ".join(escape_md(cell) for cell in header) + " |",
        "| " + " | ".join("---" for _ in header) + " |",
    ]
    for row in normalized[1:]:
        lines.append("| " + " | ".join(escape_md(cell) for cell in row) + " |")
    return "\n".join(lines)


def escape_md(text: str) -> str:
    return str(text).replace("|", "\\|")


def header_index(headers: list[str], choices: set[str]) -> int | None:
    normalized_choices = {normalize_header(choice) for choice in choices}
    for index, header in enumerate(headers):
        normalized = normalize_header(header)
        if normalized in normalized_choices:
            return index
    return None


def find_header_row(rows: list[list[str]]) -> int | None:
    for index, row in enumerate(rows):
        headers = [normalize_header(cell) for cell in row]
        has_code_or_name = header_index(headers, CODE_HEADERS) is not None or header_index(headers, NAME_HEADERS) is not None
        has_type_or_desc = header_index(headers, TYPE_HEADERS) is not None or header_index(headers, DESC_HEADERS) is not None
        if has_code_or_name and has_type_or_desc:
            return index
    return None


def mapped_header_values(header_map: dict[str, str]) -> set[str]:
    return {value for value in header_map.values() if value}


def raw_columns_for_row(headers: list[str], row: list[str]) -> dict[str, str]:
    raw: dict[str, str] = {}
    for index, header in enumerate(headers):
        header_text = cell_text(header) or f"column{index + 1}"
        key = header_text
        suffix = 2
        while key in raw:
            key = f"{header_text}#{suffix}"
            suffix += 1
        raw[key] = cell_text(row[index]) if index < len(row) else ""
    return raw


def duplicate_header_warnings(headers: list[str]) -> list[str]:
    seen: dict[str, str] = {}
    warnings: list[str] = []
    for header in headers:
        header_text = cell_text(header)
        if not header_text:
            continue
        normalized = normalize_header(header_text)
        if normalized in seen:
            warnings.append(f"conflictHeader:{seen[normalized]}|{header_text}")
        else:
            seen[normalized] = header_text
    return warnings


def build_source_location(source_base: dict[str, Any] | None, row_number: int) -> dict[str, Any]:
    location = dict(source_base or {})
    location["row"] = row_number
    return location


def required_reason(required: str, required_header: str, nullable: str, nullable_header: str) -> str:
    if required:
        if required_header:
            return f"由{required_header}={required}提供"
        return "由必填列提供"
    inferred = infer_required_from_nullable(nullable)
    if inferred:
        if nullable_header:
            return f"由{nullable_header}={nullable}推导"
        return "由允许空列推导"
    return ""


def parse_rows_to_fields(rows: list[list[str]], inherited_headers: list[str] | None = None, source_base: dict[str, Any] | None = None) -> tuple[list[Field], dict[str, str], list[str]]:
    diagnostics: list[str] = []
    if is_non_field_table(rows):
        return [], {}, ["非字段表，已跳过"]

    header_row = find_header_row(rows)
    inherited = False
    if header_row is None:
        if inherited_headers and looks_like_field_rows(rows):
            headers = inherited_headers
            data_rows = rows
            inherited = True
        else:
            return [], {}, ["未识别字段表头"]
    else:
        headers = rows[header_row]
        data_rows = rows[header_row + 1 :]

    code_i = header_index(headers, CODE_HEADERS)
    name_i = header_index(headers, NAME_HEADERS)
    type_i = header_index(headers, TYPE_HEADERS)
    length_i = header_index(headers, LENGTH_HEADERS)
    required_i = header_index(headers, REQUIRED_HEADERS)
    nullable_i = header_index(headers, NULLABLE_HEADERS)
    primary_key_i = header_index(headers, PRIMARY_KEY_HEADERS)
    default_i = header_index(headers, DEFAULT_HEADERS)
    desc_i = header_index(headers, DESC_HEADERS)

    if code_i is not None and name_i == code_i:
        name_i = alternate_name_index(headers, code_i)
    if desc_i is not None and desc_i in {code_i, name_i}:
        desc_i = alternate_description_index(headers, {index for index in [code_i, name_i] if index is not None})

    header_map = {
        "code": headers[code_i] if code_i is not None else "",
        "name": headers[name_i] if name_i is not None else "",
        "fieldType": headers[type_i] if type_i is not None else "",
        "length": headers[length_i] if length_i is not None else "",
        "required": headers[required_i] if required_i is not None else "",
        "nullable": headers[nullable_i] if nullable_i is not None else "",
        "primaryKey": headers[primary_key_i] if primary_key_i is not None else "",
        "defaultValue": headers[default_i] if default_i is not None else "",
        "description": headers[desc_i] if desc_i is not None else "",
    }

    mapped_headers = mapped_header_values(header_map)
    table_warnings = duplicate_header_warnings(headers)
    unmapped_headers = [cell_text(header) for header in headers if cell_text(header) and cell_text(header) not in mapped_headers]
    diagnostics.extend(table_warnings)

    fields: list[Field] = []
    for offset, row in enumerate(data_rows):
        if not any(cell_text(cell) for cell in row):
            continue
        get = lambda i: cell_text(row[i]) if i is not None and i < len(row) else ""
        code = get(code_i)
        if is_example_field_code(code):
            continue
        nullable = get(nullable_i)
        explicit_required = get(required_i)
        required = explicit_required or infer_required_from_nullable(nullable)
        warnings = list(table_warnings)
        warnings.extend(f"unmappedHeader:{header}" for header in unmapped_headers)
        if get(default_i):
            warnings.append("defaultValue:需要人工确认默认值语义")
        if desc_i is not None and (desc_i == code_i or desc_i == name_i):
            warnings.append("lowConfidenceDescription:说明列与代码/名称列冲突")

        row_number = (offset + 1) if inherited else (header_row or 0) + offset + 2
        field = Field(
            code=code,
            name=get(name_i),
            fieldType=get(type_i),
            length=get(length_i),
            required=required,
            nullable=nullable,
            primaryKey=get(primary_key_i),
            defaultValue=get(default_i),
            description=get(desc_i),
            requiredByMarker=code.startswith("*"),
            requiredMismatch=required_marker_mismatch(code, required),
            sourceHeaderMap=header_map,
            rawColumns=raw_columns_for_row(headers, row),
            sourceLocation=build_source_location(source_base, row_number),
            classification="mapped-field" if (code or get(name_i)) else "low-confidence",
            confidence=1.0 if (code or get(name_i)) else 0.6,
            warnings=warnings,
            requiredReason=required_reason(explicit_required, header_map["required"], nullable, header_map["nullable"]),
        )
        if field.code or field.name:
            fields.append(field)

    if not fields:
        diagnostics.append("表头已识别，但未抽取到字段行")
    return fields, header_map, diagnostics


def is_example_field_code(code: str) -> bool:
    text = cell_text(code)
    if not text:
        return False
    normalized = normalize_header(text)
    if normalized in {"\u5165\u53c2\u8868", "\u51fa\u53c2\u8868", "\u5b57\u6bb5\u540d", "\u53c2\u6570\u540d", "\u53c2\u6570\u4ee3\u7801", "\u6570\u636e\u9879\u4ee3\u7801", "\u8fd4\u56de\u76ee\u5f55"}:
        return True
    if "\u793a\u4f8b" in normalized and len(text) <= 12:
        return True
    if text.startswith(("\u6ce8\uff1a", "\u6ce8:", "{", "}", "[", "]")):
        return True
    if re.search(r"[\u4e00-\u9fff]", text) and len(text) > 20 and any(marker in text for marker in ["\u4e3a", "\u76f8\u5173", "\u4fe1\u606f", "\u8bf4\u660e"]):
        return True
    if len(text) > 20 and any(marker in text for marker in ['":"', '":', "':"]):
        return True
    return False

def clean_field_code(code: str) -> str:
    return cell_text(code).lstrip("*")


def infer_required_from_nullable(nullable: str) -> str:
    nullable_text = normalize_header(nullable)
    if nullable_text in {"\u5426", "n", "no", "false", "0", "not null"}:
        return "Y"
    if nullable_text in {"\u662f", "y", "yes", "true", "1", "null", "nullable"}:
        return "N"
    return ""

def required_marker_mismatch(code: str, required: str) -> bool:
    required_text = normalize_header(required)
    required_yes = required_text in {"y", "yes", "true", "1", "\u662f", "\u5fc5\u586b"}
    required_no = required_text in {"n", "no", "false", "0", "\u5426", "\u975e\u5fc5\u586b"}
    marker_yes = clean_field_code(code) != cell_text(code)
    if marker_yes and required_no:
        return True
    return False


def enrich_fields_for_context(fields: list[Field], context: PdfContext) -> None:
    if not context.parameterObject:
        return
    for field in fields:
        code = clean_field_code(field.code)
        if not code:
            continue
        if context.parameterObject == "headers":
            field.jsonPath = f"headers.{code}"
        elif context.parameterObject == "request":
            field.jsonPath = f"request.{code}"
        elif context.parameterObject == "data":
            field.jsonPath = f"data.{code}"
        elif context.parameterObject == "response":
            field.jsonPath = f"response.{code}"
        elif context.parameterObject == "signature":
            field.jsonPath = f"signature.{code}"
        else:
            field.jsonPath = f"data.{context.parameterObject}.{code}"
        field.jsonPathReason = f"由上下文 {context.parameterObject} 推导"

def alternate_description_index(headers: list[str], occupied: set[int]) -> int | None:
    normalized_choices = {normalize_header(choice) for choice in DESC_HEADERS}
    for index, header in enumerate(headers):
        if index in occupied:
            continue
        if normalize_header(header) in normalized_choices:
            return index
    return None


def alternate_name_index(headers: list[str], code_i: int) -> int | None:
    for choices in [{"\u63cf\u8ff0"}, {"\u4e2d\u6587\u540d", "\u540d\u79f0", "\u5b57\u6bb5\u63cf\u8ff0"}]:
        index = header_index(headers, choices)
        if index is not None and index != code_i:
            return index
    return None


def is_non_field_table(rows: list[list[str]]) -> bool:
    if not rows:
        return True
    header = {normalize_header(cell) for cell in rows[0] if cell_text(cell)}
    if {"\u65e5\u671f", "\u7248\u672c\u53f7", "\u4fee\u8ba2\u8bf4\u660e"}.issubset(header):
        return True
    if {"code", "message"}.issubset(header):
        return True
    if {"\u72b6\u6001\u7801", "\u63cf\u8ff0"}.issubset(header):
        return True
    if {"\u9519\u8bef\u7801", "\u63cf\u8ff0"}.issubset(header):
        return True
    if header in ({"\u7f16\u7801", "\u63cf\u8ff0"}, {"\u4ee3\u7801", "\u540d\u79f0"}):
        return True
    return False


def looks_like_field_rows(rows: list[list[str]]) -> bool:
    non_empty_rows = [row for row in rows if any(cell_text(cell) for cell in row)]
    if len(non_empty_rows) < 1:
        return False
    hits = 0
    for row in non_empty_rows[:8]:
        if len(row) < 2:
            continue
        code = cell_text(row[1])
        required = cell_text(row[2]).upper() if len(row) > 2 else ""
        if re.match(r"^\*?[A-Za-z_][A-Za-z0-9_]*$", code) and required in {"Y", "N", "\u662f", "\u5426", ""}:
            hits += 1
    return hits >= 1


def update_pdf_context_from_text(context: PdfContext, text: str) -> PdfContext:
    lines = [clean_section_line(line) for line in text.splitlines()]
    lines = [line for line in lines if line]
    for index, line in enumerate(lines):
        if is_toc_line(line):
            continue
        direct_interface = re.match("^(3\\.\\d+)\\s+(.+\u63a5\u53e3)\\s*$", line)
        if direct_interface:
            context.interfaceNumber = direct_interface.group(1)
            context.interfaceTitle = direct_interface.group(2).strip()
            context.subsectionNumber = ""
            context.subsectionTitle = ""
            continue

        number_only = re.match("^(3\\.\\d+)\\s*$", line)
        if number_only:
            title = nearby_interface_title(lines, index)
            if title:
                context.interfaceNumber = number_only.group(1)
                context.interfaceTitle = title
                context.subsectionNumber = ""
                context.subsectionTitle = ""
            continue

        subsection = re.match("^(3\\.\\d+\\.\\d+)\\s*(.+?)\\s*$", line)
        if subsection:
            context.subsectionNumber = subsection.group(1)
            context.subsectionTitle = subsection.group(2).strip()
            if "\u8fd4\u56de" in context.subsectionTitle:
                context.parameterObject = "response"

        object_heading = re.match("^([A-Za-z_][A-Za-z0-9_]*)\\s*\u53c2\u6570\\s*$", line)
        if object_heading:
            context.parameterObject = object_heading.group(1)
            continue

        normalized = normalize_header(line).replace(" ", "")
        if any(marker in normalized for marker in ["\u7b7e\u540d\u7b97\u6cd5", "\u53c2\u4e0e\u7b7e\u540d\u5b57\u6bb5"]):
            context.parameterObject = "signature"
            continue
        if any(marker in normalized for marker in ["\u8bf7\u6c42\u5934\u516c\u5171\u53c2\u6570", "\u516c\u5171\u8bf7\u6c42\u5934", "header\u53c2\u6570"]):
            context.parameterObject = "headers"
            continue
        if any(marker in normalized for marker in ["\u5165\u53c2\u683c\u5f0f\u8bf4\u660e", "\u5165\u53c2", "\u8bf7\u6c42\u53c2\u6570"]):
            context.parameterObject = "request"
            continue
        if any(marker in normalized for marker in ["\u54cd\u5e94\u6d88\u606f\u8bf4\u660e", "\u54cd\u5e94\u53c2\u6570", "\u8fd4\u56de\u53c2\u6570", "\u51fa\u53c2"]):
            context.parameterObject = "response"
            continue
        if "data\u53c2\u6570\u5c5e\u6027\u63cf\u8ff0\u5982\u4e0b" in normalized:
            context.parameterObject = "data"
    return context


def clean_section_line(line: str) -> str:
    text = cell_text(line)
    text = re.sub(r"\.{3,}.*$", "", text).strip()
    text = re.sub(r"\s+", " ", text)
    return text


def is_toc_line(line: str) -> bool:
    return "..." in line or bool(re.search(r"\.{3,}\s*\d+\s*$", line))


def nearby_interface_title(lines: list[str], index: int) -> str:
    for candidate in reversed(lines[max(0, index - 3) : index]):
        if candidate.endswith("\u63a5\u53e3") and not re.match(r"^\d+(\.\d+)*", candidate) and candidate != "\u63a5\u53e3\u8bf4\u660e":
            return candidate
    for candidate in lines[index + 1 : index + 4]:
        if candidate.endswith("\u63a5\u53e3") and not re.match(r"^\d+(\.\d+)*", candidate) and candidate != "\u63a5\u53e3\u8bf4\u660e":
            return candidate
    return ""


def pdf_context_label(context: PdfContext, fallback: str) -> str:
    parts: list[str] = []
    if context.interfaceNumber and context.interfaceTitle:
        parts.append(f"{context.interfaceNumber} {context.interfaceTitle}")
    elif context.interfaceTitle:
        parts.append(context.interfaceTitle)
    if context.subsectionNumber and context.subsectionTitle:
        parts.append(f"{context.subsectionNumber} {context.subsectionTitle}")
    if context.parameterObject:
        if context.parameterObject == "response":
            parts.append("response")
        elif context.parameterObject in {"headers", "request", "data", "signature"}:
            parts.append(context.parameterObject)
        else:
            parts.append(f"data.{context.parameterObject}")
    parts.append(fallback)
    return " / ".join(parts)



def copy_pdf_context(context: PdfContext) -> PdfContext:
    return PdfContext(
        interfaceNumber=context.interfaceNumber,
        interfaceTitle=context.interfaceTitle,
        subsectionNumber=context.subsectionNumber,
        subsectionTitle=context.subsectionTitle,
        parameterObject=context.parameterObject,
    )


def row_text(row: list[str]) -> str:
    return "".join(cell_text(cell) for cell in row)


def normalized_row_text(row: list[str]) -> str:
    return re.sub(r"[\s（）()]+", "", row_text(row)).lower()


def table_interface_title(rows: list[list[str]]) -> str:
    title = ""
    description = ""
    for row in rows[:6]:
        if len(row) < 2:
            continue
        label = normalize_header(row[0])
        value = cell_text(row[1])
        if not value or value == cell_text(row[0]):
            continue
        if label in {"\u540d\u79f0", "name"} and not title:
            title = value
        elif label in {"\u529f\u80fd\u63cf\u8ff0", "description"} and not description:
            description = value
    if title and description:
        return f"{title}\uff08{description}\uff09"
    return title


def row_context_marker(row: list[str]) -> str | None:
    text = normalized_row_text(row)
    if not text:
        return None
    if "接口名称" in text:
        return "interface"
    if any(marker in text for marker in ["请求参数", "请求参数加密前", "入参", "入参格式说明"]):
        return "request"
    if any(marker in text for marker in ["返回结果", "响应参数", "响应消息说明", "出参"]):
        return "response"
    return None


def row_without_section_marker(row: list[str]) -> list[str]:
    cleaned = list(row)
    for index, value in enumerate(cleaned):
        if row_context_marker([value]) is not None:
            cleaned[index] = ""
            break
    return trim_trailing_empty(cleaned)


def infer_interface_hint_from_fields(fields: list[Field]) -> str:
    for field in fields:
        if clean_field_code(field.code).lower() != "n_type":
            continue
        text = " ".join([field.name, field.description]).replace(" ", "")
        match = re.search(r"[A-Z][A-Z0-9_]{3,}", text)
        if match:
            return match.group(0)
    return ""

def parse_rows_to_context_segments(rows: list[list[str]], base_context: PdfContext, source_base: dict[str, Any] | None = None) -> list[tuple[PdfContext, list[Field], list[str], list[str] | None]]:
    segments: list[tuple[PdfContext, list[list[str]]]] = []
    current_context = copy_pdf_context(base_context)
    current_rows: list[list[str]] = []
    saw_marker = False

    def flush() -> None:
        nonlocal current_rows
        useful_rows = [row for row in current_rows if any(cell_text(cell) for cell in row)]
        if useful_rows:
            segments.append((copy_pdf_context(current_context), useful_rows))
        current_rows = []

    for row in rows:
        marker = row_context_marker(row)
        if marker == "interface":
            flush()
            saw_marker = True
            continue
        if marker in {"request", "response"}:
            flush()
            saw_marker = True
            current_context = copy_pdf_context(base_context)
            current_context.parameterObject = marker
            header_candidate = row_without_section_marker(row)
            if any(cell_text(cell) for cell in header_candidate):
                current_rows.append(header_candidate)
            continue
        current_rows.append(row)
    flush()

    if not saw_marker:
        return []

    parsed_segments: list[tuple[PdfContext, list[Field], list[str], list[str] | None]] = []
    current_interface_hint = ""
    for context, segment_rows in segments:
        fields, _header_map, diagnostics = parse_rows_to_fields(segment_rows, source_base=source_base)
        if fields:
            interface_hint = infer_interface_hint_from_fields(fields)
            if interface_hint:
                current_interface_hint = interface_hint
            if current_interface_hint:
                context.interfaceNumber = ""
                context.interfaceTitle = current_interface_hint
                context.subsectionNumber = ""
                context.subsectionTitle = ""
            enrich_fields_for_context(fields, context)
            header_row = find_header_row(segment_rows)
            inherited = segment_rows[header_row] if header_row is not None else None
            parsed_segments.append((context, fields, diagnostics, inherited))
        elif diagnostics:
            parsed_segments.append((context, [], diagnostics, None))
    return parsed_segments

def parse_xlsx(path: Path) -> tuple[str, list[View], list[str], str]:
    if has_module("openpyxl"):
        return parse_xlsx_openpyxl(path)
    return parse_xlsx_openxml(path)


def parse_xls(path: Path) -> tuple[str, list[View], list[str], str]:
    try:
        import xlrd
    except ImportError as exc:
        raise RuntimeError("XLS 解析需要 xlrd；请先运行 iris-interface-env-check.py 查看安装建议，或手动另存为 XLSX") from exc

    workbook = xlrd.open_workbook(str(path))
    markdown_parts = [f"# {path.name}", ""]
    views: list[View] = []
    diagnostics: list[str] = []

    for sheet in workbook.sheets():
        rows: list[list[str]] = []
        for row_index in range(sheet.nrows):
            row = [cell_text(sheet.cell_value(row_index, col_index)) for col_index in range(sheet.ncols)]
            row = trim_trailing_empty(row)
            if any(cell_text(cell) for cell in row):
                rows.append(row)
        markdown_parts.extend([f"## {sheet.name}", ""])
        if rows:
            markdown_parts.append(markdown_table(rows))
            markdown_parts.append("")
        fields, _header_map, field_diags = parse_rows_to_fields(rows, source_base={"documentType": "xls", "sheet": sheet.name})
        diagnostics.extend([f"{sheet.name}: {item}" for item in field_diags])
        if fields:
            views.append(View(viewCode=sheet.name, viewName=sheet.name, fields=fields))

    if not views:
        diagnostics.append("未从 XLS 中抽取到字段视图")
    return "\n".join(markdown_parts).strip() + "\n", views, diagnostics, "xls-xlrd-optional"

def has_module(module_name: str) -> bool:
    try:
        __import__(module_name)
        return True
    except ImportError:
        return False


def parse_xlsx_openpyxl(path: Path) -> tuple[str, list[View], list[str], str]:
    from openpyxl import load_workbook

    workbook = load_workbook(path, data_only=True, read_only=True)
    markdown_parts = [f"# {path.name}", ""]
    views: list[View] = []
    diagnostics: list[str] = []

    for sheet in workbook.worksheets:
        rows = [[cell_text(value) for value in row] for row in sheet.iter_rows(values_only=True)]
        rows = [trim_trailing_empty(row) for row in rows if any(cell_text(cell) for cell in row)]
        markdown_parts.extend([f"## {sheet.title}", ""])
        if rows:
            markdown_parts.append(markdown_table(rows))
            markdown_parts.append("")
        fields, _header_map, field_diags = parse_rows_to_fields(rows, source_base={"documentType": "xlsx", "sheet": sheet.title})
        diagnostics.extend([f"{sheet.title}: {item}" for item in field_diags])
        if fields:
            views.append(View(viewCode=sheet.title, viewName=sheet.title, fields=fields))

    if not views:
        diagnostics.append("\u672a\u4ece XLSX \u4e2d\u62bd\u53d6\u5230\u5b57\u6bb5\u89c6\u56fe")
    return "\n".join(markdown_parts).strip() + "\n", views, diagnostics, "xlsx-openpyxl-optional"


def parse_xlsx_openxml(path: Path) -> tuple[str, list[View], list[str], str]:
    with zipfile.ZipFile(path) as archive:
        shared_strings = read_shared_strings(archive)
        sheets = read_workbook_sheets(archive)
        rels = read_workbook_relationships(archive)
        markdown_parts = [f"# {path.name}", ""]
        views: list[View] = []
        diagnostics: list[str] = []

        for sheet_index, (sheet_name, rel_id) in enumerate(sheets, start=1):
            target = rels.get(rel_id, f"worksheets/sheet{sheet_index}.xml")
            sheet_path = "xl/" + target.lstrip("/")
            sheet_path = sheet_path.replace("xl/xl/", "xl/")
            if sheet_path not in archive.namelist():
                diagnostics.append(f"{sheet_name}: \u672a\u627e\u5230\u5de5\u4f5c\u8868 XML {sheet_path}")
                continue
            rows = read_sheet_rows(archive, sheet_path, shared_strings)
            rows = [trim_trailing_empty(row) for row in rows if any(cell_text(cell) for cell in row)]
            markdown_parts.extend([f"## {sheet_name}", ""])
            if rows:
                markdown_parts.append(markdown_table(rows))
                markdown_parts.append("")
            fields, _header_map, field_diags = parse_rows_to_fields(rows, source_base={"documentType": "xlsx", "sheet": sheet_name})
            diagnostics.extend([f"{sheet_name}: {item}" for item in field_diags])
            if fields:
                views.append(View(viewCode=sheet_name, viewName=sheet_name, fields=fields))

    if not views:
        diagnostics.append("\u672a\u4ece XLSX \u4e2d\u62bd\u53d6\u5230\u5b57\u6bb5\u89c6\u56fe")
    return "\n".join(markdown_parts).strip() + "\n", views, diagnostics, "xlsx-openxml-built-in"


def read_shared_strings(archive: zipfile.ZipFile) -> list[str]:
    if "xl/sharedStrings.xml" not in archive.namelist():
        return []
    root = ET.fromstring(archive.read("xl/sharedStrings.xml"))
    ns = {"x": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
    values: list[str] = []
    for item in root.findall("x:si", ns):
        texts = [node.text or "" for node in item.findall(".//x:t", ns)]
        values.append("".join(texts))
    return values


def read_workbook_sheets(archive: zipfile.ZipFile) -> list[tuple[str, str]]:
    root = ET.fromstring(archive.read("xl/workbook.xml"))
    ns = {"x": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
    rel_key = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
    sheets: list[tuple[str, str]] = []
    for sheet in root.findall(".//x:sheet", ns):
        sheets.append((sheet.attrib.get("name", "Sheet"), sheet.attrib.get(rel_key, "")))
    return sheets


def read_workbook_relationships(archive: zipfile.ZipFile) -> dict[str, str]:
    rel_path = "xl/_rels/workbook.xml.rels"
    if rel_path not in archive.namelist():
        return {}
    root = ET.fromstring(archive.read(rel_path))
    rels: dict[str, str] = {}
    for rel in root:
        rel_id = rel.attrib.get("Id", "")
        target = rel.attrib.get("Target", "")
        if rel_id and target:
            rels[rel_id] = target
    return rels


def read_sheet_rows(archive: zipfile.ZipFile, sheet_path: str, shared_strings: list[str]) -> list[list[str]]:
    root = ET.fromstring(archive.read(sheet_path))
    ns = {"x": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
    rows: list[list[str]] = []
    for row in root.findall(".//x:row", ns):
        values: list[str] = []
        for cell in row.findall("x:c", ns):
            col_index = column_index(cell.attrib.get("r", ""))
            while len(values) < col_index:
                values.append("")
            values.append(read_cell_value(cell, shared_strings, ns))
        rows.append(values)
    return rows


def column_index(ref: str) -> int:
    letters = re.sub(r"[^A-Z]", "", ref.upper())
    index = 0
    for letter in letters:
        index = index * 26 + (ord(letter) - ord("A") + 1)
    return max(index - 1, 0)


def read_cell_value(cell: ET.Element, shared_strings: list[str], ns: dict[str, str]) -> str:
    cell_type = cell.attrib.get("t", "")
    if cell_type == "inlineStr":
        texts = [node.text or "" for node in cell.findall(".//x:t", ns)]
        return cell_text("".join(texts))
    value_node = cell.find("x:v", ns)
    value = value_node.text if value_node is not None else ""
    if cell_type == "s" and value:
        try:
            return cell_text(shared_strings[int(value)])
        except (ValueError, IndexError):
            return ""
    return cell_text(value)

def trim_trailing_empty(row: list[str]) -> list[str]:
    result = list(row)
    while result and not cell_text(result[-1]):
        result.pop()
    return result


def parse_docx(path: Path) -> tuple[str, list[View], list[str], str]:
    try:
        import docx
    except ImportError as exc:
        raise RuntimeError("缺少 python-docx，无法解析 DOCX；请先运行 iris-interface-env-check.py 查看安装建议") from exc

    document = docx.Document(str(path))
    markdown_parts = [f"# {path.name}", ""]
    all_views: list[View] = []
    diagnostics: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            markdown_parts.extend([text, ""])

    for index, table in enumerate(document.tables, start=1):
        rows = [[cell_text(cell.text) for cell in row.cells] for row in table.rows]
        markdown_parts.extend([f"## 表格 {index}", "", markdown_table(rows), ""])

        segmented = parse_rows_to_context_segments(rows, PdfContext(interfaceTitle=table_interface_title(rows)), {"documentType": "docx", "table": index})
        segment_with_fields = [(context, fields, diags, _headers) for context, fields, diags, _headers in segmented if fields]
        if segment_with_fields:
            for segment_index, (segment_context, fields, field_diags, _headers) in enumerate(segment_with_fields, start=1):
                diagnostics.extend([f"表格 {index} Segment {segment_index}: {item}" for item in field_diags])
                segment_name = f"表格 {index} Segment {segment_index}"
                all_views.append(View(viewCode=f"table-{index}-segment-{segment_index}", viewName=pdf_context_label(segment_context, segment_name), fields=fields))
            continue

        fields, _header_map, field_diags = parse_rows_to_fields(rows, source_base={"documentType": "docx", "table": index})
        diagnostics.extend([f"表格 {index}: {item}" for item in field_diags])
        if fields:
            all_views.append(View(viewCode=f"table-{index}", viewName=f"表格 {index}", fields=fields))

    if not all_views:
        diagnostics.append("未从 DOCX 表格中抽取到字段视图")
    return "\n".join(markdown_parts).strip() + "\n", all_views, diagnostics, "docx-built-in"


def parse_pdf(path: Path) -> tuple[str, list[View], list[str], str]:
    try:
        import pdfplumber
    except ImportError as exc:
        raise RuntimeError("\u7f3a\u5c11 pdfplumber\uff0c\u65e0\u6cd5\u89e3\u6790 PDF") from exc

    markdown_parts = [f"# {path.name}", ""]
    views: list[View] = []
    diagnostics: list[str] = []
    current_field_headers: list[str] | None = None
    current_context = PdfContext()
    current_interface_hint = ""
    with pdfplumber.open(path) as pdf:
        for page_index, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                markdown_parts.extend([f"## Page {page_index}", "", text.strip(), ""])

            table_objects = sorted(page.find_tables() or [], key=lambda item: (item.bbox[1], item.bbox[0]))
            for table_index, table_object in enumerate(table_objects, start=1):
                table = table_object.extract() or []
                rows = [[cell_text(cell) for cell in row] for row in table]
                rows = [trim_trailing_empty(row) for row in rows if any(cell_text(cell) for cell in row)]
                markdown_parts.extend([f"### Page {page_index} Table {table_index}", "", markdown_table(rows), ""])

                table_context = PdfContext(
                    interfaceNumber=current_context.interfaceNumber,
                    interfaceTitle=current_context.interfaceTitle,
                    subsectionNumber=current_context.subsectionNumber,
                    subsectionTitle=current_context.subsectionTitle,
                
                    parameterObject=current_context.parameterObject,
                )
                try:
                    table_top = max(float(table_object.bbox[1]), 0.0)
                    before_table = page.crop((0, 0, page.width, table_top)).extract_text() or ""
                except Exception:
                    before_table = text
                update_pdf_context_from_text(table_context, before_table)
                if current_interface_hint:
                    table_context.interfaceNumber = ""
                    table_context.interfaceTitle = current_interface_hint
                    table_context.subsectionNumber = ""
                    table_context.subsectionTitle = ""

                table_name = f"Page {page_index} Table {table_index}"
                segmented = parse_rows_to_context_segments(rows, table_context, {"documentType": "pdf", "page": page_index, "table": table_index})
                segment_with_fields = [(context, fields, diags, inherited) for context, fields, diags, inherited in segmented if fields]
                if segment_with_fields:
                    for segment_index, (segment_context, fields, field_diags, inherited_headers) in enumerate(segment_with_fields, start=1):
                        diagnostics.extend([f"Page {page_index} Table {table_index} Segment {segment_index}: {item}" for item in field_diags])
                        if inherited_headers is not None:
                            current_field_headers = inherited_headers
                        segment_name = f"{table_name} Segment {segment_index}"
                        views.append(View(viewCode=f"p{page_index}-t{table_index}-s{segment_index}", viewName=pdf_context_label(segment_context, segment_name), fields=fields))
                        interface_hint = infer_interface_hint_from_fields(fields)
                        if interface_hint:
                            current_interface_hint = interface_hint
                    continue

                explicit_header_row = find_header_row(rows)
                fields, _header_map, field_diags = parse_rows_to_fields(rows, inherited_headers=current_field_headers, source_base={"documentType": "pdf", "page": page_index, "table": table_index})
                diagnostics.extend([f"Page {page_index} Table {table_index}: {item}" for item in field_diags])
                if fields:
                    if explicit_header_row is not None:
                        current_field_headers = rows[explicit_header_row]
                    enrich_fields_for_context(fields, table_context)
                    views.append(View(viewCode=f"p{page_index}-t{table_index}", viewName=pdf_context_label(table_context, table_name), fields=fields))
                    interface_hint = infer_interface_hint_from_fields(fields)
                    if interface_hint:
                        current_interface_hint = interface_hint

            update_pdf_context_from_text(current_context, text)

    if not views:
        diagnostics.append("\u672a\u4ece PDF \u8868\u683c\u4e2d\u62bd\u53d6\u5230\u5b57\u6bb5\u89c6\u56fe")
    return "\n".join(markdown_parts).strip() + "\n", views, diagnostics, "pdf-built-in"


def parse_doc(path: Path) -> tuple[str, list[View], list[str], str]:
    converted = try_convert_doc(path)
    if converted is not None:
        return parse_document(converted)

    markitdown_markdown = try_markitdown(path)
    if markitdown_markdown:
        diagnostics = ["MarkItDown \u4ec5\u751f\u6210 Markdown\uff1bDOC \u7ed3\u6784\u5316\u5b57\u6bb5\u62bd\u53d6\u9700\u8981\u8f6c\u4e3a DOCX \u540e\u590d\u6838"]
        return markitdown_markdown, [], diagnostics, "markitdown-optional"

    raise RuntimeError("DOC \u89e3\u6790\u9700\u8981 LibreOffice/Pandoc \u8f6c DOCX\uff1bMarkItDown \u4ec5\u80fd\u4f5c\u4e3a Markdown \u964d\u7ea7\uff0c\u8bf7\u5b89\u88c5\u8f6c\u6362\u5668\u6216\u624b\u52a8\u8f6c\u4e3a DOCX")


def try_markitdown(path: Path) -> str | None:
    try:
        from markitdown import MarkItDown
    except ImportError:
        return None

    try:
        result = MarkItDown().convert(str(path))
        text = getattr(result, "text_content", "") or ""
    except Exception:
        return None

    if not text.strip():
        return None
    return f"# {path.name}\n\n" + text.strip() + "\n"


def try_convert_doc(path: Path) -> Path | None:
    temp_dir = path.parent / f".{path.stem}-converted"
    temp_dir.mkdir(exist_ok=True)

    libreoffice = shutil.which("soffice") or shutil.which("libreoffice")
    if libreoffice:
        subprocess.run(
            [libreoffice, "--headless", "--convert-to", "docx", "--outdir", str(temp_dir), str(path)],
            check=False,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        candidate = temp_dir / f"{path.stem}.docx"
        if candidate.exists():
            return candidate

    pandoc = shutil.which("pandoc")
    if pandoc:
        candidate = temp_dir / f"{path.stem}.docx"
        subprocess.run([pandoc, str(path), "-o", str(candidate)], check=False, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        if candidate.exists():
            return candidate

    return None


def parse_document(path: Path) -> tuple[str, list[View], list[str], str]:
    suffix = path.suffix.lower()
    if suffix == ".xlsx":
        return parse_xlsx(path)
    if suffix == ".xls":
        return parse_xls(path)
    if suffix == ".docx":
        return parse_docx(path)
    if suffix == ".pdf":
        return parse_pdf(path)
    if suffix == ".doc":
        return parse_doc(path)
    raise RuntimeError(f"不支持的文件类型: {suffix}")


def artifact_json(source: Path, doc_name: str, views: list[View], diagnostics: list[str], converter: str) -> dict[str, Any]:
    view_dicts = []
    total_fields = 0
    for view in views:
        fields = [asdict(field) for field in view.fields]
        total_fields += len(fields)
        view_dicts.append({"viewCode": view.viewCode, "viewName": view.viewName, "fields": fields})

    return {
        "schemaVersion": "iris-interface-doc-ingest/v2",
        "sourceFile": str(source),
        "documentName": doc_name,
        "converter": converter,
        "viewCount": len(view_dicts),
        "totalFields": total_fields,
        "views": view_dicts,
        "diagnostics": diagnostics,
    }


def trace_summary(field: Field) -> str:
    items = [field.requiredReason, field.jsonPathReason]
    items.extend(field.warnings)
    return "; ".join(item for item in items if item)


def fields_markdown(views: list[View]) -> str:
    lines = ["# Fields", ""]
    for view in views:
        lines.extend([f"## {view.viewName}", "", "| JSON路径 | 字段代码 | 字段名称 | 类型 | 长度 | 必填 | 允许空 | 主键 | 默认值 | 必填标记 | 备注 | 追溯提示 |", "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |"])
        for field in view.fields:
            required_marker = "*" if field.requiredByMarker else ""
            if field.requiredMismatch:
                required_marker += " mismatch"
            lines.append(
                "| "
                + " | ".join(
                    escape_md(value)
                    for value in [field.jsonPath, field.code, field.name, field.fieldType, field.length, field.required, field.nullable, field.primaryKey, field.defaultValue, required_marker, field.description, trace_summary(field)]
                )
                + " |"
            )
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def diagnostics_markdown(parsed: dict[str, Any]) -> str:
    fields = [field for view in parsed.get("views", []) for field in view.get("fields", [])]
    field_warning_count = sum(len(field.get("warnings") or []) for field in fields)
    low_confidence_count = sum(1 for field in fields if float(field.get("confidence") or 0) < 1.0)
    inferred_required_count = sum(1 for field in fields if "推导" in str(field.get("requiredReason") or ""))
    lines = [
        "# Diagnostics",
        "",
        f"- converter: `{parsed['converter']}`",
        f"- views: `{parsed['viewCount']}`",
        f"- fields: `{parsed['totalFields']}`",
        f"- fieldWarnings: `{field_warning_count}`",
        f"- lowConfidenceFields: `{low_confidence_count}`",
        f"- inferredRequiredFields: `{inferred_required_count}`",
        "",
        "## Notes",
        "",
    ]
    diagnostics = parsed.get("diagnostics") or []
    if diagnostics:
        lines.extend(f"- {item}" for item in diagnostics)
    else:
        lines.append("- 未发现解析级错误；字段语义匹配仍需后续诊断。")
    return "\n".join(lines).strip() + "\n"


def write_artifacts(source: Path, project_root: Path, output_root: str) -> dict[str, Path]:
    markdown, views, diagnostics, converter = parse_document(source)
    doc_name = slugify(source)
    out_dir = project_root / output_root / doc_name
    out_dir.mkdir(parents=True, exist_ok=True)

    parsed = artifact_json(source, doc_name, views, diagnostics, converter)
    paths = {
        "source": out_dir / "source.md",
        "parsed": out_dir / "parsed.json",
        "fields": out_dir / "fields.md",
        "diagnostics": out_dir / "diagnostics.md",
    }
    paths["source"].write_text(markdown, encoding="utf-8")
    paths["parsed"].write_text(json.dumps(parsed, ensure_ascii=False, indent=2), encoding="utf-8")
    paths["fields"].write_text(fields_markdown(views), encoding="utf-8")
    paths["diagnostics"].write_text(diagnostics_markdown(parsed), encoding="utf-8")
    return paths


def parse_args(argv: Iterable[str]) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Parse interface document artifacts to disk.")
    parser.add_argument("--file", required=True, help="Source DOCX/PDF/XLSX/DOC file")
    parser.add_argument("--project-root", default=".", help="Target project root")
    parser.add_argument("--output-root", default="docs/output/iris-interface", help="Output root relative to project root")
    return parser.parse_args(argv)


def main(argv: Iterable[str]) -> int:
    args = parse_args(argv)
    source = Path(args.file).resolve()
    project_root = Path(args.project_root).resolve()
    if not source.exists():
        print(f"ERROR: source file not found: {source}", file=sys.stderr)
        return 2

    try:
        paths = write_artifacts(source, project_root, args.output_root)
    except Exception as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1

    print("iris-interface-doc-ingest completed")
    for key in ["source", "parsed", "fields", "diagnostics"]:
        print(f"{key}: {paths[key]}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
